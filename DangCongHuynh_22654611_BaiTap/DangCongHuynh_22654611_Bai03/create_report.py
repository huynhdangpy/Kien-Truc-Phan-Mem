"""
Script tạo file Word báo cáo
Yêu cầu: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """Set background color for cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_style(doc, text, level=1):
    """Add heading with custom style"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def create_report():
    """Create Word document"""
    doc = Document()
    
    # ============================================
    # TITLE PAGE
    # ============================================
    title = doc.add_heading('BÀI 3: SO SÁNH REST, gRPC VÀ GraphQL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Xây dựng Hệ thống Backend Hiện Đại', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Course info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Tác giả: ').bold = True
    info.add_run('Đặng Công Huyền (22654611)\n')
    info.add_run('Môn học: ').bold = True
    info.add_run('Kiến Trúc Phần Mềm\n')
    info.add_run('Ngày hoàn thành: ').bold = True
    info.add_run('02/02/2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        '1. Khái Niệm',
        '2. Ví Dụ Thực Tế',
        '3. So Sánh Chi Tiết',
        '4. Phân Tích Chuyên Sâu',
        '5. Kết Luận & Đề Xuất'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================
    # 1. CONCEPTS
    # ============================================
    doc.add_heading('1. KHÁI NIỆM', level=1)
    
    # REST
    doc.add_heading('1.1 REST (Representational State Transfer)', level=2)
    doc.add_paragraph('REST là kiến trúc phần mềm dựa trên HTTP, sử dụng HTTP methods (GET, POST, PUT, DELETE) để thực hiện các phép toán trên tài nguyên.')
    
    doc.add_heading('Đặc điểm chính:', level=3)
    rest_features = [
        'Dựa trên HTTP verbs: GET, POST, PUT, DELETE',
        'Tài nguyên được xác định bằng URL paths',
        'Định dạng dữ liệu: JSON, XML',
        'Stateless: mỗi request độc lập',
        'Easy to understand và widely adopted'
    ]
    for feature in rest_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Ưu điểm:', level=3)
    rest_pros = [
        'Dễ hiểu và dễ học',
        'Hỗ trợ tốt trên browser và tools',
        'Caching dễ dàng',
        'Scalable và simple'
    ]
    for pro in rest_pros:
        doc.add_paragraph(pro, style='List Bullet')
    
    doc.add_heading('Nhược điểm:', level=3)
    rest_cons = [
        'Over-fetching: lấy dữ liệu không cần thiết',
        'Under-fetching: cần nhiều requests',
        'Versioning API khó quản lý'
    ]
    for con in rest_cons:
        doc.add_paragraph(con, style='List Bullet')
    
    # gRPC
    doc.add_heading('1.2 gRPC (gRPC Remote Procedure Call)', level=2)
    doc.add_paragraph('gRPC là framework hiện đại dựa trên HTTP/2 cho việc gọi hàm từ xa, sử dụng Protocol Buffers (protobuf) để serialize dữ liệu.')
    
    doc.add_heading('Đặc điểm chính:', level=3)
    grpc_features = [
        'Dựa trên HTTP/2 multiplexing',
        'Định dạng dữ liệu: Protocol Buffers (binary)',
        'Hỗ trợ 4 loại streaming',
        'Schema-first (phải định nghĩa .proto)',
        'Hiệu năng cao, độ trễ thấp'
    ]
    for feature in grpc_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Ưu điểm:', level=3)
    grpc_pros = [
        'Hiệu năng rất cao (binary protocol)',
        'Hỗ trợ real-time streaming',
        'Định dạng nhỏ gọn (tiết kiệm bandwidth)',
        'Type-safe (Protocol Buffers)',
        'Tuyệt vời cho microservices'
    ]
    for pro in grpc_pros:
        doc.add_paragraph(pro, style='List Bullet')
    
    doc.add_heading('Nhược điểm:', level=3)
    grpc_cons = [
        'Không hỗ trợ browser trực tiếp',
        'Khó debug (binary format)',
        'Yêu cầu học Protocol Buffers',
        'Ecosystem nhỏ hơn REST'
    ]
    for con in grpc_cons:
        doc.add_paragraph(con, style='List Bullet')
    
    # GraphQL
    doc.add_heading('1.3 GraphQL (Query Language)', level=2)
    doc.add_paragraph('GraphQL là ngôn ngữ truy vấn cho APIs, cho phép client yêu cầu chính xác những dữ liệu họ cần.')
    
    doc.add_heading('Đặc điểm chính:', level=3)
    graphql_features = [
        'Ngôn ngữ truy vấn, không phải giao thức',
        'Một endpoint duy nhất: POST /graphql',
        'Định dạng dữ liệu: JSON',
        'Client-driven: client yêu cầu cái gì, server trả về cái đó',
        'Strong typing (schema)'
    ]
    for feature in graphql_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Ưu điểm:', level=3)
    graphql_pros = [
        'Lấy chính xác dữ liệu cần thiết',
        'Type system mạnh mẽ',
        'Single endpoint: dễ quản lý',
        'Tuyệt vời cho mobile',
        'Dễ versioning'
    ]
    for pro in graphql_pros:
        doc.add_paragraph(pro, style='List Bullet')
    
    doc.add_heading('Nhược điểm:', level=3)
    graphql_cons = [
        'Phức tạp hơn để học',
        'Performance: complex queries có thể chậm',
        'Caching khó hơn (POST-based)',
        'N+1 query problem'
    ]
    for con in graphql_cons:
        doc.add_paragraph(con, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================
    # 2. COMPARISON TABLE
    # ============================================
    doc.add_heading('2. SO SÁNH CHI TIẾT', level=1)
    
    # Create comparison table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Tiêu Chí', 'REST', 'gRPC', 'GraphQL']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], 'E7E6E6')
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    comparison_data = [
        ['Giao Thức', 'HTTP/1.1', 'HTTP/2', 'HTTP/1.1'],
        ['Định Dạng Dữ Liệu', 'JSON, XML', 'Protobuf (binary)', 'JSON'],
        ['Phương Thức Truy Vấn', 'Multiple URLs', 'Method calls', 'Single endpoint'],
        ['Caching', 'HTTP (dễ)', 'Caching khó', 'POST-based (khó)'],
        ['Hiệu Năng', 'Trung bình', 'Rất cao', 'Trung bình'],
        ['Độ Trễ', '~50-100ms', '~5-10ms', '~50-100ms'],
        ['Payload Size', '~1-3KB', '~100-500B', '~1-3KB'],
        ['Browser Support', '✅ Gốc', '❌ Cần gRPC-web', '✅ Gốc'],
        ['Streaming', '❌ WebSocket', '✅ Native', '✅ WebSocket'],
        ['Type Safety', '❌', '✅ Protobuf', '✅ GraphQL Schema'],
        ['Learning Curve', '⭐ Dễ', '⭐⭐⭐ Khó', '⭐⭐ Trung bình'],
    ]
    
    for row_data in comparison_data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
    
    doc.add_page_break()
    
    # ============================================
    # 3. CODE EXAMPLES
    # ============================================
    doc.add_heading('3. VÍ DỤ CODE', level=1)
    
    doc.add_heading('3.1 REST API - Express.js', level=2)
    doc.add_paragraph('Tạo user mới:')
    rest_code = '''app.post('/api/users', (req, res) => {
  const { name, email } = req.body;
  if (!name || !email) {
    return res.status(400).json({ error: 'Missing fields' });
  }
  const newUser = { id: users.length + 1, name, email };
  users.push(newUser);
  res.status(201).json({ success: true, data: newUser });
});'''
    doc.add_paragraph(rest_code, style='List Number')
    
    doc.add_heading('3.2 gRPC API - Protocol Buffer', level=2)
    doc.add_paragraph('Proto định nghĩa:')
    proto_code = '''service UserService {
  rpc GetUser(GetUserRequest) returns (GetUserResponse);
  rpc CreateUser(CreateUserRequest) returns (CreateUserResponse);
  rpc StreamUserPosts(GetUserPostsRequest) returns (stream Post);
}'''
    doc.add_paragraph(proto_code, style='List Number')
    
    doc.add_heading('3.3 GraphQL API - Apollo Server', level=2)
    doc.add_paragraph('Query schema:')
    graphql_code = '''type Query {
  user(id: Int!): User
  users: [User!]!
}

type Mutation {
  createUser(name: String!, email: String!): User!
}'''
    doc.add_paragraph(graphql_code, style='List Number')
    
    doc.add_page_break()
    
    # ============================================
    # 4. USE CASES
    # ============================================
    doc.add_heading('4. PHÂN TÍCH CHUYÊN SÂU', level=1)
    
    doc.add_heading('4.1 REST Chạy Tốt Nhất Khi?', level=2)
    doc.add_paragraph('Trường hợp sử dụng tối ưu:')
    rest_cases = [
        'Public APIs & Web Services (Facebook, GitHub)',
        'CRUD Operations Đơn Giản',
        'HTTP Caching Quan Trọng (CDN)',
        'Stateless Services (Easy scaling)'
    ]
    for case in rest_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    doc.add_heading('4.2 gRPC Vượt Trội Khi?', level=2)
    doc.add_paragraph('Trường hợp sử dụng tối ưu:')
    grpc_cases = [
        'Microservices Communication (Netflix, Google)',
        'Real-time Streaming Data (IoT, Trading)',
        'High-Performance Systems (latency <10ms)',
        'Internal Infrastructure (performance > UX)'
    ]
    for case in grpc_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    doc.add_heading('4.3 Khi Nào Dùng GraphQL?', level=2)
    doc.add_paragraph('GraphQL giải quyết các vấn đề:')
    graphql_cases = [
        'Over-fetching Problem: lấy chỉ field cần',
        'Under-fetching Problem: 1 query thay vì 3-4',
        'Mobile Apps: Bandwidth optimization 30-50%',
        'Complex Nested Data: Social media, CMS',
        'Rapid Frontend Development: Type safety'
    ]
    for case in graphql_cases:
        doc.add_paragraph(case, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================
    # 5. RECOMMENDATIONS
    # ============================================
    doc.add_heading('5. KẾT LUẬN & ĐỀ XUẤT', level=1)
    
    # Create recommendations table
    rec_table = doc.add_table(rows=1, cols=2)
    rec_table.style = 'Light Grid Accent 1'
    
    rec_headers = rec_table.rows[0].cells
    rec_headers[0].text = 'Loại Hệ Thống'
    rec_headers[1].text = 'Công Nghệ Đề Xuất'
    rec_headers[0].paragraphs[0].runs[0].bold = True
    rec_headers[1].paragraphs[0].runs[0].bold = True
    set_cell_background(rec_headers[0], 'E7E6E6')
    set_cell_background(rec_headers[1], 'E7E6E6')
    
    recommendations = [
        ['Microservices', 'gRPC (chính) + REST Gateway'],
        ['Mobile Apps', 'GraphQL (chính) + REST Fallback'],
        ['Web Apps', 'REST (simple) hoặc GraphQL (complex)'],
        ['Hệ thống Nội Bộ HP', 'gRPC + Event-Driven Architecture']
    ]
    
    for rec in recommendations:
        rec_cells = rec_table.add_row().cells
        rec_cells[0].text = rec[0]
        rec_cells[1].text = rec[1]
    
    # Final conclusions
    doc.add_heading('Khuyến Nghị Cuối Cùng', level=2)
    conclusions = [
        'REST: Chọn khi công khai API, caching quan trọng, team mới',
        'gRPC: Chọn khi performance <10ms, microservices, internal systems',
        'GraphQL: Chọn khi mobile app, bandwidth optimization, complex queries',
        'Hybrid: Sử dụng kết hợp: REST Gateway + gRPC Services + GraphQL for Mobile'
    ]
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================
    # PROJECT STRUCTURE
    # ============================================
    doc.add_heading('6. CẤU TRÚC DỰ ÁN', level=1)
    
    structure_text = '''DangCongHuynh_22654611_Bai03/
├── BAI_3_SO_SANH_REST_gRPC_GraphQL.md
├── SO_SANH_CODE_IMPLEMENTATION.md
├── REST-API/
│   ├── server.js
│   ├── client.js
│   ├── package.json
│   └── README.md
├── gRPC-API/
│   ├── server.js
│   ├── client.js
│   ├── api.proto
│   ├── package.json
│   └── README.md
└── GraphQL-API/
    ├── server.js
    ├── client.js
    ├── package.json
    └── README.md'''
    
    doc.add_paragraph(structure_text, style='List Number')
    
    doc.add_heading('Hướng Dẫn Chạy', level=2)
    
    doc.add_heading('REST API', level=3)
    doc.add_paragraph('cd REST-API && npm install && npm start', style='List Bullet')
    doc.add_paragraph('Test: node client.js', style='List Bullet')
    
    doc.add_heading('gRPC API', level=3)
    doc.add_paragraph('cd gRPC-API && npm install && npm start', style='List Bullet')
    doc.add_paragraph('Test: npm run client', style='List Bullet')
    
    doc.add_heading('GraphQL API', level=3)
    doc.add_paragraph('cd GraphQL-API && npm install && npm start', style='List Bullet')
    doc.add_paragraph('Truy cập: http://localhost:4000', style='List Bullet')
    
    # Save document
    doc.save('BAI_3_So_Sanh_REST_gRPC_GraphQL_Report.docx')
    print('✓ File Word đã tạo: BAI_3_So_Sanh_REST_gRPC_GraphQL_Report.docx')

if __name__ == '__main__':
    create_report()
